import os
from flask import Flask, render_template, request, redirect, url_for, session, send_file
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
app.secret_key = "your_secret_key"

UPLOAD_FOLDER = "generated_papers"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Replace with your valid API key
genai.configure(api_key="AIzaSyB5RRdbSHe9K2FYRTrcGpiIhEjw-myna1Q")

users = {"admin": "password123"}

@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        if username in users and users[username] == password:
            session["username"] = username
            return redirect(url_for("dashboard"))
        return render_template("login.html", error="Invalid credentials")
    return render_template("login.html")

@app.route("/dashboard")
def dashboard():
    if "username" not in session:
        return redirect(url_for("login"))
    return render_template("dashboard.html")

@app.route("/logout")
def logout():
    session.pop("username", None)
    return redirect(url_for("login"))

def generate_questions(syllabus_text):
    prompt = f"""
    Generate a B.Tech VI Semester (R20) Computer Networks Question Paper with the following strict format:

    PART A: 10 questions (2 marks each)
    Format:
    1. Question text here (CO: CO1, BT: BTL1)

    PART B: 5 units with 2 questions each (10 marks)
    Format:
    UNIT - I
    1. Long question text here (CO: CO2, BT: BTL3)
    (OR)
    2. Alternative question text here (CO: CO2, BT: BTL4)

    Follow the exact format above. Include only questions and unit titles. No introductions or explanations.

    Syllabus:
    {syllabus_text}
    """
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(prompt)
    return response.text if response else ""

@app.route("/generate", methods=["GET", "POST"])
def generate():
    if "username" not in session:
        return redirect(url_for("login"))

    if request.method == "POST":
        exam_title = request.form.get("exam_title", "B.TECH - VI SEMESTER (R20)")
        subject = request.form.get("subject", "COMPUTER NETWORKS")
        syllabus_text = request.form["syllabus_text"]
        qp_code = request.form.get("qp_code", "20CSE363")
        question_paper_text = generate_questions(syllabus_text)

        if not question_paper_text:
            return render_template("generate.html", error="Failed to generate questions. Try again later.")

        doc = Document()

        # Set page margins
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Header
        p1 = doc.add_paragraph("HALL TICKET NO.: ___________")
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Q. P. Code: {qp_code}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(exam_title).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("REGULAR / SUPPLEMENTARY EXAMINATIONS - JUN - 2024").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(subject).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("(Common to CSE, CSM, CAI and CSD)").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        time_marks_paragraph = doc.add_paragraph("Time: 3 Hours\t\t\t\t\t\t\t\tMax. Marks: 70")
        time_marks_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("---------------------------------------------------------------------").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # PART A
        p_part_a = doc.add_paragraph("PART – A", style='Heading 2')
        p_part_a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("Answer ALL questions. Each question carries 2 marks.\n").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Process the generated question paper text
        part_a = ""
        part_b = ""
        if "PART B" in question_paper_text:
            parts = question_paper_text.split("PART B")
            part_a = parts[0].replace("PART A", "").strip()
            part_b = "PART B\n" + parts[1].strip()
        else:
            lines = question_paper_text.split("\n")
            part_a_lines = [line for line in lines if "CO:" in line and "BT:" in line]
            part_b_lines = [line for line in lines if "UNIT" in line]
            part_a = "\n".join(part_a_lines)
            part_b = "\n".join(part_b_lines)

        # Add Part A questions
        question_counter = 1  # To keep track of question numbers across both parts
        for line in part_a.split("\n"):
            if line.strip() and any(char.isdigit() for char in line):
                parts = line.strip().split("(")
                question = parts[0].strip()
                co_bt = parts[1].replace(")", "").split(",") if len(parts) > 1 else []

                co = co_bt[0].strip() if len(co_bt) > 0 else "CO: -"
                bt = co_bt[1].strip() if len(co_bt) > 1 else "BT: -"

                qno = question.split(".")[0]
                qtext = ".".join(question.split(".")[1:]).strip()

                doc.add_paragraph(f"{question_counter}. {qtext} ({co}, {bt})")
                question_counter += 1

        # PART B
        if part_b:
            p_part_b = doc.add_paragraph("\nPART – B", style='Heading 2')
            p_part_b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("Answer ONE question from each UNIT – Each question carries 10 marks.\n").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            units = part_b.strip().split("UNIT")
            for unit in units[1:]:
                lines = unit.strip().split("\n")
                unit_title = "UNIT - " + lines[0].strip()
                doc.add_paragraph(unit_title, style='Heading 3').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                q_data = []
                for line in lines[1:]:
                    if line.strip() and "." in line:
                        qn = line.split(".")[0].strip()
                        question = ".".join(line.split(".")[1:]).strip()

                        co_bt = []
                        if "(" in question:
                            parts = question.split("(")
                            qtext = parts[0].strip()
                            meta = "".join(parts[1:]).replace(")", "")
                            co_bt = meta.split(",")
                        else:
                            qtext = question
                        co = co_bt[0].strip() if len(co_bt) > 0 else "CO: -"
                        bt = co_bt[1].strip() if len(co_bt) > 1 else "BT: -"

                        q_data.append(f"{question_counter}. {qtext} ({co}, {bt})")
                        question_counter += 1

                # Add questions to the document with "OR" inserted in the middle of the question text
                if len(q_data) > 1:
                    doc.add_paragraph(q_data[0])  # First question with CO and BT
                    doc.add_paragraph("OR")  # Insert "OR" between the two questions
                    doc.add_paragraph(q_data[1])  # Second question with CO and BT
                elif len(q_data) == 1:
                    doc.add_paragraph(q_data[0])  # If only one question is present, add it

        filename = os.path.join(app.config["UPLOAD_FOLDER"], "Formatted_Question_Paper.docx")
        doc.save(filename)
        return send_file(filename, as_attachment=True)

    return render_template("generate.html")

if __name__ == "__main__":
    app.run(debug=True)
